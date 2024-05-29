import openpyxl
from docx import Document
from collections import defaultdict

def read_excel(file_path):
    workbook = openpyxl.load_workbook(file_path)
    data = defaultdict(list)

    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        print(f"正在读取工作表: {sheet_name}")

        for row in sheet.iter_rows(values_only=True):
            # 根据工作表名来处理不同的列数
            if sheet_name == "选择题":
                if len(row) < 7:
                    continue
                question_type, question, correct_answer, colA, colB, colC, colD = row[:7]
                question_data = {
                    "题型": question_type,
                    "题干": question,
                    "A": colA,
                    "B": colB,
                    "C": colC,
                    "D": colD,
                    "正确答案": correct_answer
                }
            elif sheet_name in ["简答题", "填空题", "判断题"]:
                if len(row) < 3:
                    continue
                question_type, question, correct_answer = row[:3]
                question_data = {
                    "题型": question_type,
                    "题干": question,
                    "正确答案": correct_answer
                }
            else:
                continue  # 如果遇到不认识的工作表，则跳过

            data[sheet_name].append(question_data)

    return data

def write_to_word(data, output_path):
    doc = Document()
    question_order = ["选择题", "判断题", "填空题", "简答题"]

    for question_type in question_order:
        if question_type in data:
            doc.add_heading(question_type, level=1)

            for i, question_data in enumerate(data[question_type], start=1):
                doc.add_paragraph(f"{i}.")
                for key, value in question_data.items():
                    doc.add_paragraph(f"{key}: {value}")
                doc.add_paragraph()

            doc.add_page_break()
            print(f"{question_type} 导出完毕")

    doc.save(output_path)

def main(excel_path, word_path):
    data = read_excel(excel_path)
    write_to_word(data, word_path)
    print(f"数据已从{excel_path}拷贝到{word_path}")

if __name__ == "__main__":
    excel_path = 'UBase.xlsx'
    word_path = 'UBase3.docx'
    main(excel_path, word_path)
